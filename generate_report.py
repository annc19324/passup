# generate_report.py
# Tạo file báo cáo thực tập môn Hệ thống thông tin tích hợp - Lê Thiên An

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# =============================================
# CẤU HÌNH TRANG
# =============================================
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(3)
section.right_margin = Cm(2)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# =============================================
# HÀM TIỆN ÍCH
# =============================================
def set_font(run, bold=False, size=13, color=None, italic=False):
    run.font.name = 'Times New Roman'
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure correct font for Vietnamese
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1, size=14, bold=True, center=False, upper=False, space_before=12, space_after=6):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size * 1.5)
    r = p.add_run(text.upper() if upper else text)
    set_font(r, bold=bold, size=size)
    return p

def add_paragraph(doc, text, indent=True, size=13, space_before=3, space_after=3, line_spacing=20, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    r = p.add_run(text)
    set_font(r, size=size, italic=italic)
    return p

def add_bullet(doc, text, size=13, line_spacing=18, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    p.paragraph_format.left_indent = Cm(1.5)
    # Thêm ký hiệu bullet thủ công
    if bold_prefix:
        r1 = p.add_run(f'• {bold_prefix}: ')
        set_font(r1, bold=True, size=size)
        r2 = p.add_run(text)
        set_font(r2, size=size)
    else:
        r = p.add_run(f'• {text}')
        set_font(r, size=size)
    return p

def add_numbered(doc, num, text, size=13, line_spacing=18, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    p.paragraph_format.left_indent = Cm(1.5)
    if bold_prefix:
        r1 = p.add_run(f'{num}. {bold_prefix}: ')
        set_font(r1, bold=True, size=size)
        r2 = p.add_run(text)
        set_font(r2, size=size)
    else:
        r = p.add_run(f'{num}. {text}')
        set_font(r, size=size)
    return p

def add_page_break(doc):
    doc.add_page_break()

def add_sub_bullet(doc, text, size=13, line_spacing=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    p.paragraph_format.left_indent = Cm(2.5)
    r = p.add_run(f'- {text}')
    set_font(r, size=size)
    return p

# =============================================
# TRANG BÌA
# =============================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
r = p.add_run('BỘ CÔNG THƯƠNG')
set_font(r, bold=True, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
r = p.add_run('TRƯỜNG ĐẠI HỌC ĐIỆN LỰC')
set_font(r, bold=True, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
r = p.add_run('KHOA CÔNG NGHỆ THÔNG TIN')
set_font(r, bold=True, size=13)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
r = p.add_run('THỰC TẬP HỆ THỐNG THÔNG TIN TÍCH HỢP')
set_font(r, bold=True, size=14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
r = p.add_run('XÂY DỰNG NỀN TẢNG MUA BÁN ĐỒ CŨ TRỰC TUYẾN PASSUP')
set_font(r, bold=True, size=16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TÍCH HỢP HỆ THỐNG QUẢN LÝ HÌNH ẢNH VÀ TRUYỀN THÔNG THỜI GIAN THỰC')
set_font(r, bold=True, size=14)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Bảng thông tin sinh viên
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'

def fill_bìa_row(row_idx, label, value):
    row = table.rows[row_idx]
    cell_label = row.cells[0]
    cell_value = row.cells[1]

    p_label = cell_label.paragraphs[0]
    r_label = p_label.add_run(label)
    set_font(r_label, bold=True, size=13)

    p_value = cell_value.paragraphs[0]
    r_value = p_value.add_run(value)
    set_font(r_value, size=13)

fill_bìa_row(0, 'Giảng viên hướng dẫn', ': [Tên Giảng viên]')
fill_bìa_row(1, 'Sinh viên thực hiện', ': Lê Thiên An')
fill_bìa_row(2, 'Mã sinh viên', ': 22810310030')
fill_bìa_row(3, 'Ngành', ': Công nghệ thông tin')
fill_bìa_row(4, 'Chuyên ngành', ': Công nghệ phần mềm')
fill_bìa_row(5, 'Lớp', ': D17CNPM1')
fill_bìa_row(6, 'Khóa', ': 2022 - 2027')
fill_bìa_row(7, 'Năm học', ': 2025 - 2026')

for row in table.rows:
    for cell in row.cells:
        cell._element.attrib.pop('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top', None)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Hà Nội, tháng 03 năm 2026')
set_font(r, size=13, italic=True)

add_page_break(doc)

# =============================================
# MỤC LỤC (tự tạo)
# =============================================
add_heading(doc, 'MỤC LỤC', level=1, size=14, center=True, space_before=0)

toc_items = [
    ('LỜI CẢM ƠN', ''),
    ('LỜI NÓI ĐẦU', ''),
    ('ĐỀ CƯƠNG THỰC TẬP', ''),
    ('CHƯƠNG 1: ĐẶT VẤN ĐỀ CẦN GIẢI QUYẾT', ''),
    ('    1.1. Lý do chọn đề tài', ''),
    ('    1.2. Mục tiêu của hệ thống', ''),
    ('    1.3. Giải pháp tổng quan', ''),
    ('CHƯƠNG 2: CHI TIẾT GIẢI PHÁP KỸ THUẬT CÔNG NGHỆ', ''),
    ('    2.1. Cơ sở kiến thức liên quan', ''),
    ('    2.2. Chi tiết giải pháp tích hợp', ''),
    ('    2.3. Kiến trúc hệ thống', ''),
    ('CHƯƠNG 3: TRIỂN KHAI GIẢI PHÁP', ''),
    ('    3.1. Phân tích thiết kế hệ thống', ''),
    ('    3.2. Cơ sở dữ liệu vật lý', ''),
    ('    3.3. Kết quả thực nghiệm', ''),
    ('KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', ''),
    ('TÀI LIỆU THAM KHẢO', ''),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(20)
    r = p.add_run(item)
    set_font(r, bold=('CHƯƠNG' in item or item.strip() in ['LỜI CẢM ƠN', 'LỜI NÓI ĐẦU', 'ĐỀ CƯƠNG THỰC TẬP', 'KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', 'TÀI LIỆU THAM KHẢO', 'MỤC LỤC']), size=12)

add_page_break(doc)

# =============================================
# LỜI CẢM ƠN
# =============================================
add_heading(doc, 'LỜI CẢM ƠN', level=1, size=14, center=True, space_before=0)

add_paragraph(doc, 'Trong suốt quá trình thực hiện báo cáo thực tập môn Thực tập Hệ thống thông tin tích hợp, em đã nhận được rất nhiều sự hỗ trợ, giúp đỡ quý báu từ thầy cô và bạn bè.')
add_paragraph(doc, 'Em xin gửi lời cảm ơn chân thành và sâu sắc đến quý thầy cô Khoa Công nghệ Thông tin, Trường Đại học Điện Lực, đặc biệt là giảng viên hướng dẫn đã tận tình chỉ bảo, định hướng và truyền đạt kiến thức trong suốt thời gian thực tập.')
add_paragraph(doc, 'Xin chân thành cảm ơn các anh, chị và các bạn học lớp D17CNPM1 đã luôn động viên, chia sẻ kinh nghiệm và hỗ trợ trong quá trình học tập và nghiên cứu.')
add_paragraph(doc, 'Cuối cùng, em xin bày tỏ lòng biết ơn sâu sắc tới gia đình đã luôn là chỗ dựa tinh thần và tạo điều kiện tốt nhất để em hoàn thành báo cáo này.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(12)
r = p.add_run('Hà Nội, ngày 29 tháng 03 năm 2026')
set_font(r, size=13, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Sinh viên thực hiện')
set_font(r, bold=True, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(36)
r = p.add_run('Lê Thiên An')
set_font(r, bold=True, size=13)

add_page_break(doc)

# =============================================
# LỜI NÓI ĐẦU
# =============================================
add_heading(doc, 'LỜI NÓI ĐẦU', level=1, size=14, center=True, space_before=0)

add_paragraph(doc, 'Trong bối cảnh công nghệ thông tin phát triển mạnh mẽ, các ứng dụng thương mại điện tử ngày càng trở nên phổ biến và đóng vai trò quan trọng trong cuộc sống hàng ngày. Đặc biệt, với đối tượng sinh viên, nhu cầu mua bán, trao đổi đồ dùng cũ ngày càng tăng cao do áp lực tài chính cũng như nhu cầu luân chuyển tài sản linh hoạt.')
add_paragraph(doc, 'Đề tài "Xây dựng nền tảng mua bán đồ cũ trực tuyến PassUp – Tích hợp hệ thống quản lý hình ảnh và truyền thông thời gian thực" được thực hiện nhằm giải quyết bài toán thực tế này. Đây là một hệ thống thông tin tích hợp, kết hợp nhiều công nghệ và dịch vụ bên thứ ba như Cloudinary (quản lý hình ảnh) và Socket.io (truyền thông thời gian thực) để tạo nên một sản phẩm hoàn chỉnh.')
add_paragraph(doc, 'Nội dung báo cáo bao gồm 03 chương:')

add_paragraph(doc, 'Chương 1: Đặt vấn đề cần giải quyết – Trình bày lý do, mục tiêu và giải pháp tổng quan của đề tài.', indent=False)
add_paragraph(doc, 'Chương 2: Chi tiết giải pháp kỹ thuật công nghệ – Phân tích các công nghệ và cơ sở lý thuyết được sử dụng, trọng tâm là các điểm tích hợp hệ thống bên thứ ba.', indent=False)
add_paragraph(doc, 'Chương 3: Triển khai giải pháp – Trình bày thiết kế hệ thống, cơ sở dữ liệu và kết quả cài đặt, demo.', indent=False)

add_paragraph(doc, 'Do kiến thức và kinh nghiệm còn hạn chế, báo cáo khó tránh khỏi những thiếu sót. Em rất mong nhận được sự góp ý của thầy cô để nội dung báo cáo được hoàn thiện hơn.')

add_page_break(doc)

# =============================================
# ĐỀ CƯƠNG THỰC TẬP
# =============================================
add_heading(doc, 'ĐỀ CƯƠNG THỰC TẬP MÔN HỌC', level=1, size=14, center=True, space_before=0)
add_heading(doc, 'THỰC TẬP HỆ THỐNG THÔNG TIN TÍCH HỢP', level=1, size=13, center=True)

doc.add_paragraph()

def add_decuong_row(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(20)
    r1 = p.add_run(f'{label}')
    set_font(r1, bold=bold_label, size=13)
    r2 = p.add_run(f'  {value}')
    set_font(r2, size=13)

add_decuong_row(doc, '1. Tên đề tài:', 'Xây dựng nền tảng mua bán đồ cũ trực tuyến PassUp – Tích hợp hệ thống quản lý hình ảnh và truyền thông thời gian thực.')
add_decuong_row(doc, '2. Sinh viên thực hiện:', '')
add_decuong_row(doc, '   Họ và tên:', 'Lê Thiên An', bold_label=False)
add_decuong_row(doc, '   Số điện thoại:', '[Số điện thoại]', bold_label=False)
add_decuong_row(doc, '   Vị trí thực tập:', 'Sinh viên – Nhà trường', bold_label=False)
add_decuong_row(doc, '3. Giảng viên hướng dẫn:', '')
add_decuong_row(doc, '   Họ và tên:', '[Tên Giảng viên]', bold_label=False)
add_decuong_row(doc, '   Số điện thoại:', '[Số điện thoại]', bold_label=False)
add_decuong_row(doc, '   Đơn vị công tác:', 'Khoa Công nghệ Thông tin, Trường Đại học Điện Lực.', bold_label=False)

add_paragraph(doc, '', indent=False)
p = doc.add_paragraph()
r = p.add_run('4. Mô tả tóm tắt đề tài:')
set_font(r, bold=True, size=13)

add_paragraph(doc, 'PassUp là một nền tảng web tích hợp (Integrated Information System) được thiết kế để hỗ trợ việc mua bán, trao đổi đồ cũ trong cộng đồng sinh viên. Hệ thống kết hợp RESTful API (Node.js/Express), cơ sở dữ liệu quan hệ (PostgreSQL thông qua Prisma ORM), giao diện người dùng hiện đại (React + TypeScript), và hai dịch vụ bên thứ ba quan trọng: nền tảng quản lý hình ảnh Cloudinary và thư viện Socket.io để xây dựng tính năng chat thời gian thực.')

p = doc.add_paragraph()
r = p.add_run('5. Nội dung báo cáo thực tập:')
set_font(r, bold=True, size=13)

add_paragraph(doc, 'Chương 1. Đặt vấn đề cần giải quyết: Nêu vấn đề thực tiễn, mục tiêu và phương hướng giải quyết tổng quan.', indent=False)
add_paragraph(doc, 'Chương 2. Chi tiết giải pháp kỹ thuật công nghệ: Trình bày cơ sở lý thuyết và các điểm tích hợp kỹ thuật (Cloudinary, Socket.io, Prisma).', indent=False)
add_paragraph(doc, 'Chương 3. Triển khai giải pháp: Cài đặt, trình bày thiết kế hệ thống và demo kết quả.', indent=False)
add_paragraph(doc, 'Kết luận: Tóm tắt kết quả đạt được, hạn chế và hướng phát triển.', indent=False)

add_page_break(doc)

# =============================================
# CHƯƠNG 1
# =============================================
add_heading(doc, 'CHƯƠNG 1', size=14, center=True, space_before=0)
add_heading(doc, 'ĐẶT VẤN ĐỀ CẦN GIẢI QUYẾT', size=14, center=True, space_before=0)

add_heading(doc, '1.1. Lý do chọn đề tài', size=13, bold=True, space_before=12)

add_paragraph(doc, 'Trong môi trường đại học, sinh viên thường có nhu cầu rất lớn trong việc thanh lý và tìm kiếm các đồ vật đã qua sử dụng như giáo trình, tài liệu học tập, đồ gia dụng nhỏ, thiết bị điện tử cũ với giá cả hợp lý. Tuy nhiên, việc kết nối giữa người có nhu cầu bán và người có nhu cầu mua trong cùng khu vực hiện nay vẫn còn nhiều khó khăn.')
add_paragraph(doc, 'Các nền tảng thương mại điện tử lớn hiện nay như Shopee, Lazada hay Chợ Tốt dù rất phổ biến nhưng chưa thực sự tối ưu cho nhu cầu giao dịch nội bộ sinh viên – nơi mà yếu tố địa lý gần (cùng khu vực, cùng trường), sự tin tưởng lẫn nhau và khả năng giao tiếp trực tiếp đóng vai trò then chốt để chốt giao dịch thành công.')
add_paragraph(doc, 'Xuất phát từ thực tế đó, đề tài PassUp được đề xuất nhằm xây dựng một hệ thống thông tin tích hợp, chuyên biệt cho cộng đồng sinh viên, giúp kết nối người mua và người bán một cách nhanh chóng, minh bạch và an toàn trong phạm vi khu vực địa lý cụ thể.')

add_heading(doc, '1.2. Mục tiêu của hệ thống', size=13, bold=True, space_before=12)

add_bullet(doc, 'Xây dựng một nền tảng web hiện đại, giao diện thân thiện, tương thích tốt với thiết bị di động.')
add_bullet(doc, 'Tích hợp dịch vụ lưu trữ và xử lý hình ảnh bên thứ ba (Cloudinary) để tối ưu hóa tài nguyên đa phương tiện của hệ thống.')
add_bullet(doc, 'Tích hợp giao thức WebSocket (Socket.io) để xây dựng hệ thống chat thời gian thực giúp người mua và người bán trao đổi trực tiếp.')
add_bullet(doc, 'Cung cấp hệ thống lọc theo địa lý (Tỉnh/Thành, Quận/Huyện, Phường/Xã) để hỗ trợ giao dịch trực tiếp.')
add_bullet(doc, 'Xây dựng hệ thống quản lý đơn hàng và thanh toán nội bộ đơn giản.')

add_heading(doc, '1.3. Tổng quan về giải pháp', size=13, bold=True, space_before=12)

add_paragraph(doc, 'Hệ thống PassUp được thiết kế theo mô hình Client-Server hiện đại với kiến trúc phân tách rõ ràng giữa giao diện người dùng (Frontend) và logic xử lý nghiệp vụ (Backend). Các thành phần chính bao gồm:')
add_bullet(doc, 'Frontend: ReactJS + TypeScript + TailwindCSS, được đóng gói và phục vụ bởi Vite.', bold_prefix=None)
add_bullet(doc, 'Backend: Node.js (Express framework) + Prisma ORM + PostgreSQL.', bold_prefix=None)
add_bullet(doc, 'Tích hợp thứ nhất: Cloudinary API – Quản lý và phân phối hình ảnh sản phẩm.', bold_prefix=None)
add_bullet(doc, 'Tích hợp thứ hai: Socket.io – Truyền thông thời gian thực (Real-time Communication).', bold_prefix=None)

add_page_break(doc)

# =============================================
# CHƯƠNG 2
# =============================================
add_heading(doc, 'CHƯƠNG 2', size=14, center=True, space_before=0)
add_heading(doc, 'CHI TIẾT GIẢI PHÁP KỸ THUẬT CÔNG NGHỆ', size=14, center=True, space_before=0)

add_heading(doc, '2.1. Cơ sở kiến thức liên quan', size=13, bold=True, space_before=12)

add_heading(doc, '2.1.1. Hệ thống thông tin tích hợp (Integrated Information System)', size=12, bold=True, space_before=8)
add_paragraph(doc, 'Hệ thống thông tin tích hợp là tập hợp các hệ thống con, dịch vụ và ứng dụng phần mềm khác nhau được kết nối và phối hợp hoạt động với nhau để phục vụ một mục tiêu nghiệp vụ chung. Điểm mấu chốt của một hệ thống tích hợp là khả năng trao đổi dữ liệu liền mạch giữa các thành phần, ngay cả khi chúng được xây dựng bởi các nhà cung cấp khác nhau.')
add_paragraph(doc, 'Trong dự án PassUp, các "hệ thống bên thứ ba" được tích hợp bao gồm: dịch vụ lưu trữ đám mây Cloudinary và thư viện giao tiếp thời gian thực Socket.io. Chúng được xem là "blackbox" – hệ thống chỉ cần biết cách gọi API/giao thức của chúng mà không cần quan tâm đến cơ chế nội bộ bên trong.')

add_heading(doc, '2.1.2. RESTful API', size=12, bold=True, space_before=8)
add_paragraph(doc, 'REST (Representational State Transfer) là một kiến trúc phần mềm dùng để xây dựng các dịch vụ web. Dựa trên giao thức HTTP, RESTful API sử dụng các phương thức GET, POST, PUT, DELETE để thực hiện các thao tác CRUD (Create, Read, Update, Delete) trên tài nguyên. Đây là phương thức giao tiếp chính giữa Frontend và Backend trong PassUp.')

add_heading(doc, '2.1.3. WebSocket & Socket.io', size=12, bold=True, space_before=8)
add_paragraph(doc, 'WebSocket là giao thức truyền thông hai chiều (full-duplex) trên nền TCP, cho phép Server và Client giao tiếp liên tục mà không cần thiết lập kết nối mới với mỗi yêu cầu. Khác với mô hình Request-Response của HTTP truyền thống, WebSocket tạo ra một kênh truyền tải liên tục (persistent connection).')
add_paragraph(doc, 'Socket.io là một thư viện JavaScript được xây dựng trên nền WebSocket, bổ sung thêm các cơ chế như: tự động kết nối lại, phân phòng (Room), phát sóng (Broadcast) và hỗ trợ fallback HTTP long-polling khi WebSocket không khả dụng.')

add_heading(doc, '2.1.4. Cloudinary – Dịch vụ quản lý hình ảnh đám mây', size=12, bold=True, space_before=8)
add_paragraph(doc, 'Cloudinary là một nền tảng quản lý tài nguyên đa phương tiện (DAM – Digital Asset Management) trên đám mây. Nó cung cấp API để tải lên, lưu trữ, biến đổi và phân phối hình ảnh/video thông qua mạng CDN (Content Delivery Network) toàn cầu. Trong PassUp, Cloudinary đóng vai trò là "hệ thống quản lý hình ảnh tích hợp" – giúp giảm tải hoàn toàn gánh nặng lưu trữ file khỏi server chính.')

add_heading(doc, '2.2. Chi tiết giải pháp tích hợp trong PassUp', size=13, bold=True, space_before=12)

add_heading(doc, '2.2.1. Tích hợp Cloudinary API', size=12, bold=True, space_before=8)

add_bullet(doc, 'Người dùng chọn ảnh trên giao diện, Frontend gửi qua HTTP POST dưới dạng multipart/form-data.', bold_prefix='Luồng hoạt động (1)')
add_sub_bullet(doc, 'Trên Server, thư viện Multer xử lý file upload và giữ dữ liệu ảnh trong bộ nhớ (memoryStorage).')
add_sub_bullet(doc, 'Server đọc buffer của file, chuyển thành base64 Data URI và gọi cloudinary.uploader.upload().')
add_sub_bullet(doc, 'Cloudinary nhận ảnh, tự động: resize về 800x800px (crop=fill), nén và lưu trong CDN.')
add_sub_bullet(doc, 'Cloudinary trả về secure_url – URL HTTPS của ảnh đã xử lý.')
add_sub_bullet(doc, 'Server lưu URL này vào cột images[] của bảng Product trong PostgreSQL.')

add_paragraph(doc, 'Đây là đoạn code minh họa luồng tích hợp Cloudinary trong product.service.ts:', indent=True)

code_p = doc.add_paragraph()
code_p.paragraph_format.left_indent = Cm(1.5)
code_p.paragraph_format.space_before = Pt(4)
code_p.paragraph_format.space_after = Pt(4)
code_run = code_p.add_run(
    'const uploadResponse = await cloudinary.uploader.upload(dataURI, {\n'
    '    folder: "passup_products",\n'
    '    resource_type: "image",\n'
    '    width: 800, height: 800,\n'
    '    crop: "fill", gravity: "center"\n'
    '});\n'
    'imageUrls.push(uploadResponse.secure_url);'
)
set_font(code_run, size=10, italic=True)

add_bullet(doc, 'Giảm tải hoàn toàn việc lưu trữ file cho Server chính, không cần cấu hình disk storage trên VPS.', bold_prefix='Lợi ích (2)')
add_sub_bullet(doc, 'Tự động tối ưu hóa kích thước ảnh giúp tốc độ tải trang nhanh hơn đáng kể.')
add_sub_bullet(doc, 'Phân phối qua CDN toàn cầu, người dùng ở bất kỳ đâu đều nhận được ảnh nhanh.')

add_heading(doc, '2.2.2. Tích hợp Socket.io (Chat thời gian thực)', size=12, bold=True, space_before=8)

add_bullet(doc, 'Khi vào trang Chat, trình duyệt thiết lập kết nối WebSocket với Server và join vào Room mang ID của người dùng (ví dụ: user_5).', bold_prefix='Luồng hoạt động (1)')
add_sub_bullet(doc, 'Khi người dùng gửi tin nhắn, Client gọi API /api/chat/send. Server lưu tin nhắn vào Database.')
add_sub_bullet(doc, 'Ngay sau đó, Server gọi io.to("user_{receiverId}").emit("new_message", data) để đẩy tin nhắn tới người nhận.')
add_sub_bullet(doc, 'Nếu người nhận đang mở trang Chat, tin nhắn sẽ hiện ngay lập tức mà không cần F5.')

add_bullet(doc, 'Trải nghiệm người dùng mượt mà như Messenger, giúp thương lượng và chốt đơn hàng nhanh hơn.', bold_prefix='Lợi ích (2)')
add_sub_bullet(doc, 'Hệ thống Room của Socket.io đảm bảo tin nhắn chỉ đến đúng người nhận.')

add_heading(doc, '2.2.3. Tích hợp Prisma ORM & PostgreSQL', size=12, bold=True, space_before=8)

add_paragraph(doc, 'Prisma là một ORM (Object-Relational Mapper) thế hệ mới cho Node.js, cho phép định nghĩa schema cơ sở dữ liệu bằng ngôn ngữ khai báo (Prisma Schema Language) và tự động sinh ra TypeScript client với đầy đủ type-safety. Trong PassUp, Prisma quản lý toàn bộ các model dữ liệu: User, Product, Order, Conversation, SentMessage, Rating, Report, Notification.')

add_heading(doc, '2.3. Kiến trúc hệ thống PassUp', size=13, bold=True, space_before=12)

add_paragraph(doc, 'Kiến trúc tổng thể của PassUp được thiết kế theo mô hình 3 tầng (Three-Tier Architecture):')
add_bullet(doc, 'Tầng Trình bày (Presentation Layer): ReactJS + TypeScript + TailwindCSS, chạy trên trình duyệt người dùng.', bold_prefix='Tầng 1')
add_bullet(doc, 'Tầng Logic nghiệp vụ (Business Logic Layer): Node.js + Express, xử lý yêu cầu từ client, gọi tích hợp dịch vụ bên thứ ba và tương tác với Database.', bold_prefix='Tầng 2')
add_bullet(doc, 'Tầng Dữ liệu (Data Layer): PostgreSQL (Neon cloud) – lưu trữ toàn bộ dữ liệu quan hệ. Cloudinary – lưu trữ tài nguyên đa phương tiện.', bold_prefix='Tầng 3')

add_page_break(doc)

# =============================================
# CHƯƠNG 3
# =============================================
add_heading(doc, 'CHƯƠNG 3', size=14, center=True, space_before=0)
add_heading(doc, 'TRIỂN KHAI GIẢI PHÁP', size=14, center=True, space_before=0)

add_heading(doc, '3.1. Phân tích và thiết kế hệ thống', size=13, bold=True, space_before=12)

add_heading(doc, '3.1.1. Xác định yêu cầu và tác nhân', size=12, bold=True, space_before=8)

add_paragraph(doc, 'Hệ thống PassUp phục vụ 3 nhóm tác nhân chính:')
add_bullet(doc, 'Có thể duyệt, tìm kiếm và xem chi tiết sản phẩm mà không cần đăng nhập.', bold_prefix='Khách vãng lai (Guest)')
add_bullet(doc, 'Đăng ký tài khoản, đăng bán sản phẩm, đặt hàng, chat, quản lý đơn hàng, yêu thích sản phẩm, đánh giá người bán, báo cáo vi phạm.', bold_prefix='Người dùng đã đăng nhập (User)')
add_bullet(doc, 'Quản lý toàn bộ người dùng, danh mục sản phẩm, xử lý báo cáo vi phạm, quản lý rút tiền và cấu hình hệ thống.', bold_prefix='Quản trị viên (Admin)')

add_heading(doc, '3.1.2. Biểu đồ Use Case tổng quát', size=12, bold=True, space_before=8)

add_paragraph(doc, 'Use Case tổng quát bao gồm các nhóm chức năng chính:')
add_bullet(doc, 'Xem sản phẩm, Tìm kiếm, Lọc theo danh mục và địa điểm.')
add_bullet(doc, 'Đăng ký, Đăng nhập, Quản lý hồ sơ cá nhân.')
add_bullet(doc, 'Đăng bán sản phẩm (upload ảnh qua Cloudinary), Quản lý tin đăng.')
add_bullet(doc, 'Chat realtime với người bán/người mua (qua Socket.io).')
add_bullet(doc, 'Đặt hàng và theo dõi đơn hàng.')
add_bullet(doc, 'Đánh giá (Rating) sau giao dịch thành công.')
add_bullet(doc, 'Wishlist – Lưu sản phẩm yêu thích.')
add_bullet(doc, 'Báo cáo vi phạm sản phẩm hoặc người dùng.')

add_heading(doc, '3.1.3. Biểu đồ trình tự – Luồng Đăng bán tích hợp Cloudinary', size=12, bold=True, space_before=8)

steps = [
    ('Bước 1', 'Người dùng điền thông tin sản phẩm và chọn ảnh trên giao diện Sell.tsx (Frontend).'),
    ('Bước 2', 'Nhấn "Đăng Bán Ngay" → Frontend tạo FormData và gọi POST /api/products với header Content-Type: multipart/form-data.'),
    ('Bước 3', 'Server nhận request → Multer xử lý file, giữ buffer trong RAM.'),
    ('Bước 4', 'product.service.ts kiểm tra giới hạn đăng tin theo gói (FREE/PRO).'),
    ('Bước 5', 'Với mỗi ảnh: chuyển buffer → base64 → gọi Cloudinary API → nhận secure_url.'),
    ('Bước 6', 'Prisma tạo bản ghi Product mới trong PostgreSQL với mảng imageUrls.'),
    ('Bước 7', 'Server trả về 201 Created → Frontend chuyển hướng về trang chủ.'),
]
for i, (step, desc) in enumerate(steps, 1):
    add_numbered(doc, i, desc, bold_prefix=step)

add_heading(doc, '3.2. Cơ sở dữ liệu vật lý', size=13, bold=True, space_before=12)

add_paragraph(doc, 'Cơ sở dữ liệu của PassUp được quản lý bởi PostgreSQL, truy cập thông qua Prisma ORM. Các bảng chính:')

db_tables = [
    ('users', 'Lưu thông tin người dùng: email, password (bcrypt), fullName, địa chỉ, số dư (balance), gói thành viên (subscriptionType), vai trò (USER/ADMIN).'),
    ('products', 'Lưu thông tin sản phẩm: tiêu đề, mô tả, giá, mảng URL ảnh (từ Cloudinary), địa chỉ, trạng thái (AVAILABLE/RESERVED/SOLD), slug SEO, cờ isHighlight, isPromoted.'),
    ('orders', 'Quản lý đơn hàng: người mua, sản phẩm, trạng thái (PENDING/CONFIRMED/SHIPPING/COMPLETED/CANCELLED), địa chỉ giao hàng.'),
    ('conversations', 'Quản lý cuộc hội thoại giữa 2 người dùng về một sản phẩm cụ thể.'),
    ('messages', 'Lưu nội dung tin nhắn: text, ảnh (nếu có), người gửi, trạng thái isRead.'),
    ('ratings', 'Đánh giá sau giao dịch: sao (1-5), bình luận, liên kết với Order.'),
    ('reports', 'Báo cáo vi phạm: loại (PRODUCT/USER), lý do, trạng thái xử lý.'),
    ('notifications', 'Thông báo hệ thống: tiêu đề, nội dung, loại, trạng thái đã đọc.'),
    ('withdrawals', 'Yêu cầu rút tiền từ người bán: số tiền, trạng thái (PENDING/APPROVED/REJECTED).'),
    ('system_settings', 'Cấu hình hệ thống (key-value), quản lý bởi Admin.'),
]

for table_name, desc in db_tables:
    add_bullet(doc, desc, bold_prefix=table_name)

add_heading(doc, '3.3. Kết quả xây dựng phần mềm', size=13, bold=True, space_before=12)

add_heading(doc, '3.3.1. Tính năng đã hoàn thiện', size=12, bold=True, space_before=8)

features = [
    'Xác thực người dùng: Đăng ký, Đăng nhập với JWT, bảo mật mật khẩu bằng bcrypt.',
    'Đăng bán sản phẩm: Upload tối đa 5 ảnh/tin, tự động tối ưu qua Cloudinary. Hỗ trợ 2 hình thức giao dịch: Tự đến lấy & Giao hàng.',
    'Trang chủ: Hiển thị sản phẩm dạng Grid. Bộ lọc thông minh theo Danh mục và Khu vực (Tỉnh/Quận/Phường). Hỗ trợ đánh dấu tin nổi bật (isHighlight).',
    'Hệ thống Chat Real-time: Nhắn tin tức thì qua Socket.io. Hỗ trợ gửi ảnh trong hội thoại.',
    'Hệ thống Đơn hàng: Quy trình PENDING → CONFIRMED → SHIPPING → COMPLETED. Tự động chuyển trạng thái sản phẩm sang RESERVED/SOLD. Thông báo tức thì cho người liên quan.',
    'Đánh giá & Nhận xét: Sau khi đơn hàng hoàn tất, người mua có thể đánh giá người bán (1-5 sao).',
    'Wishlist: Lưu sản phẩm yêu thích, xem lại bất cứ lúc nào.',
    'Hồ sơ công khai người bán: Hiển thị thống kê sản phẩm, điểm đánh giá trung bình.',
    'Báo cáo vi phạm: Báo cáo sản phẩm hoặc người dùng.',
    'Hệ thống thương mại hóa: Các gói thành viên FREE/PRO với giới hạn đăng tin. Lượt đẩy tin (Push), tin nổi bật (Highlight).',
    'Trang Admin Dashboard: Quản lý người dùng, danh mục, báo cáo vi phạm, yêu cầu rút tiền, cấu hình hệ thống.',
    'Tối ưu SEO: Slug URL thân thiện (vd: /san-pham/iphone-13-128gb-1748293929).',
    'Responsive Design: Giao diện tương thích tốt với Mobile, có thanh điều hướng bottom navigation bar.',
]

for i, feat in enumerate(features, 1):
    add_numbered(doc, i, feat)

add_heading(doc, '3.3.2. Môi trường và cách chạy hệ thống', size=12, bold=True, space_before=8)

add_bullet(doc, 'Node.js v18+, npm, Python 3.x.', bold_prefix='Yêu cầu môi trường')
add_numbered(doc, 1, 'Clone source code từ repository về máy local.', bold_prefix='Bước')
add_numbered(doc, 2, 'Cấu hình file .env trong thư mục server: DATABASE_URL (PostgreSQL), CLOUDINARY_*, JWT_SECRET.', bold_prefix='Bước')
add_numbered(doc, 3, 'Cài đặt dependencies: npm install (trong cả server/ và client/).', bold_prefix='Bước')
add_numbered(doc, 4, 'Đẩy schema lên Database: cd server && npx prisma db push.', bold_prefix='Bước')
add_numbered(doc, 5, 'Khởi động Backend: npm run dev (cổng 3000).', bold_prefix='Bước')
add_numbered(doc, 6, 'Khởi động Frontend: npm run dev (cổng 5173).', bold_prefix='Bước')

add_page_break(doc)

# =============================================
# KẾT LUẬN
# =============================================
add_heading(doc, 'KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', size=14, center=True, space_before=0)

add_heading(doc, '1. Những kết quả đạt được', size=13, bold=True, space_before=8)
add_paragraph(doc, 'Qua quá trình thực hiện đề tài, nhóm đã đạt được những kết quả sau:')
add_bullet(doc, 'Nắm vững quy trình thiết kế và xây dựng một hệ thống thông tin tích hợp hoàn chỉnh, từ phân tích yêu cầu, thiết kế CSDL đến triển khai và kiểm thử.')
add_bullet(doc, 'Triển khai thành công việc tích hợp dịch vụ bên thứ ba Cloudinary vào luồng xử lý ảnh, giúp hệ thống nhẹ hơn và ảnh được phân phối nhanh hơn thông qua CDN.')
add_bullet(doc, 'Xây dựng thành công hệ thống Chat Real-time sử dụng Socket.io, mang lại trải nghiệm người dùng mượt mà và trực quan.')
add_bullet(doc, 'Hoàn thiện nền tảng PassUp với đầy đủ các chức năng cốt lõi từ Đăng nhập, Đăng tin, Chat, Đặt hàng cho đến Admin Dashboard.')
add_bullet(doc, 'Áp dụng các best practices hiện đại: JWT Auth, bcrypt, Prisma type-safe, RESTful API conventions, SEO-friendly slug.')

add_heading(doc, '2. Những hạn chế', size=13, bold=True, space_before=8)
add_bullet(doc, 'Hệ thống thanh toán hiện tại chỉ là giả lập (mô phỏng), chưa tích hợp cổng thanh toán điện tử thực tế như Momo, VNPay.')
add_bullet(doc, 'Chưa tích hợp dịch vụ vận chuyển thực tế (GHTK, GHN) để tính phí ship tự động theo địa chỉ.')
add_bullet(doc, 'Hệ thống tìm kiếm chưa sử dụng Full-text search nâng cao (Elasticsearch), với tập dữ liệu lớn hiệu suất có thể giảm.')
add_bullet(doc, 'Chưa có ứng dụng mobile native (iOS/Android), thông báo push notification chưa được triển khai đầy đủ.')

add_heading(doc, '3. Hướng phát triển trong tương lai', size=13, bold=True, space_before=8)
add_bullet(doc, 'Tích hợp cổng thanh toán điện tử (Momo, VNPay, ZaloPay) để đảm bảo an toàn cho giao dịch online.')
add_bullet(doc, 'Tích hợp đơn vị vận chuyển (GHN, GHTK) để hỗ trợ giao hàng toàn quốc, tự động tính phí ship.')
add_bullet(doc, 'Phát triển ứng dụng di động sử dụng React Native, kế thừa backend API hiện có.')
add_bullet(doc, 'Ứng dụng AI/ML để gợi ý sản phẩm cá nhân hóa dựa trên lịch sử duyệt và hành vi người dùng.')
add_bullet(doc, 'Nâng cấp search engine sử dụng Elasticsearch để hỗ trợ tìm kiếm nâng cao, tìm kiếm gần đúng (fuzzy search).')

add_page_break(doc)

# =============================================
# TÀI LIỆU THAM KHẢO
# =============================================
add_heading(doc, 'TÀI LIỆU THAM KHẢO', size=14, center=True, space_before=0)

refs = [
    ('React Team (2024).', 'React Documentation – The library for web and native user interfaces. Truy cập ngày 20/03/2026, từ <https://react.dev/>'),
    ('Prisma (2024).', 'Prisma ORM Documentation – Next-generation Node.js and TypeScript ORM. Truy cập ngày 20/03/2026, từ <https://www.prisma.io/docs>'),
    ('Socket.io (2024).', 'Socket.IO Documentation – Bidirectional and low-latency communication. Truy cập ngày 20/03/2026, từ <https://socket.io/docs/v4/>'),
    ('Cloudinary (2024).', 'Cloudinary API Reference – Cloud-based image and video management. Truy cập ngày 20/03/2026, từ <https://cloudinary.com/documentation>'),
    ('MDN Web Docs (2024).', 'HTTP Overview – Hypertext Transfer Protocol. Truy cập ngày 15/03/2026, từ <https://developer.mozilla.org/en-US/docs/Web/HTTP/Overview>'),
    ('Express.js (2024).', 'Express - Fast, unopinionated, minimalist web framework for Node.js. Truy cập ngày 15/03/2026, từ <https://expressjs.com/>'),
    ('Multer (2024).', 'Multer – Node.js middleware for handling multipart/form-data. Truy cập ngày 18/03/2026, từ <https://github.com/expressjs/multer>'),
]

for i, (author, title) in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    r1 = p.add_run(f'[{i}] {author} ')
    set_font(r1, bold=True, size=13)
    r2 = p.add_run(title)
    set_font(r2, size=13)

# =============================================
# LƯU FILE
# =============================================
output_path = r'd:\demo\PassUp\BAO_CAO_THUC_TAP_Le_Thien_An_22810310030.docx'
doc.save(output_path)
print(f"✅ Đã tạo file báo cáo thành công: {output_path}")
